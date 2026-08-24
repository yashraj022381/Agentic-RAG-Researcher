import os
import re
import json
import docx
import pytesseract
import concurrent.futures
from pathlib import Path
from typing import Optional, Tuple
from .base import BaseTool, ToolResult
from pdf2image import convert_from_path
from pypdf import PdfReader
from utils.paths import DOCS_DIR
                

SUPPORTED = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"}

# Safety ceiling on extracted TEXT length — not the primary size-reduction
# mechanism (that happens downstream in utils/document_excerpt.py, which
# searches keywords across the full text). This just protects against
# truly pathological inputs.
MAX_CONTENT_CHARS = 6000

# Bounds how many PDF pages get processed in one call. A 150MB+ PDF can
# have thousands of pages; extracting all of them page-by-page in a single
# synchronous call can take minutes and looks indistinguishable from a
# hang. Capping this means large files fail fast and clearly, with an
# honest message, instead of silently appearing to "not work."
MAX_PDF_PAGES = 400

# Files larger than this on disk get a clear upfront message rather than
# attempting extraction at all. This is deliberately generous — most large
# PDFs are large because of embedded images, not text — but a hard ceiling
# is still needed since there's no way to know the actual text yield
# without reading it, and reading a 150MB file synchronously is not safe
# to attempt unconditionally.
MAX_FILE_SIZE_MB = 100

_GENERIC_STOPWORDS = {
    "document", "documents", "file", "files", "content", "contents",
    "page", "pages", "section", "sections", "chapter", "chapters",
    "table", "tables", "text", "format", "data", "information",
    "extract", "search", "construct", "uploaded", "current",
}


class DocumentReaderTool(BaseTool):

    def __init__(self, docs_folder: str = DOCS_DIR):
        self.docs_folder = Path(docs_folder)
        self.docs_folder.mkdir(parents=True, exist_ok=True)

    @property
    def name(self) -> str:
        return "document_reader"

    @property
    def description(self) -> str:
        return (
            "Read and extract content from uploaded documents (PDF, DOCX, TXT). "
            "Best when you need detail from a specific known file or report."
            "Supports scanned PDFs via OCR."
        )

    def run(
        self,
        query: str,
        context: Optional[str] = None,
        file_path: Optional[str] = None,
    ) -> ToolResult:
        
        q = query.lower().strip()

        available_files = [
            f for f in self.docs_folder.rglob("*")
            if f.is_file() and f.suffix.lower() in SUPPORTED
        ]

        if not available_files:
            return ToolResult(
                content=(
                    f"No documents found in '{self.docs_folder.resolve()}'. "
                    f"Please add PDF, DOCX, or TXT files to that folder."
                ),
                source="Document: folder empty",
                confidence=0.0,
            )

        matched = None

        if file_path:
            candidate = Path(file_path)
            if candidate.is_file() and candidate in available_files:
                matched = candidate

        if matched is None:
            matched = self._find_best_match(q, available_files)

        if not matched:
            names = [f.name for f in available_files]
            return ToolResult(
                content=(
                    f"No document matching '{query}' found.\n"
                    f"Available files: {', '.join(names)}"
                ),
                source="Document: not found",
                confidence=0.1,
            )

        # Explicit, honest size gate — instead of attempting extraction on
        # anything of any size and letting it silently hang or fail.
        size_mb = matched.stat().st_size / (1024 * 1024)
        if size_mb > MAX_FILE_SIZE_MB:
            return ToolResult(
                content=(
                    f"'{matched.name}' is {size_mb:.0f} MB, which exceeds the "
                    f"{MAX_FILE_SIZE_MB} MB processing limit for direct text "
                    f"extraction. Very large files need chunked/embedding-based "
                    f"retrieval rather than full-text extraction — this isn't "
                    f"supported yet. Consider splitting out the relevant "
                    f"section into a smaller file."
                ),
                source=f"Document: {matched.name} (too large)",
                confidence=0.1,
            )

        try:
            content, truncation_note = self._read_file(matched)
            if not content.strip():
                return ToolResult(
                    content=f"File '{matched.name}' appears to be empty.",
                    source=f"Document: {matched.name}",
                    confidence=0.2,
                )

            truncated = content[:MAX_CONTENT_CHARS]
            if len(content) > MAX_CONTENT_CHARS:
                truncated += f"\n\n[... {len(content) - MAX_CONTENT_CHARS} more characters truncated ...]"
            if truncation_note:
                truncated += f"\n\n[NOTE: {truncation_note}]"

            query_words = set(w for w in re.findall(r'\w+', query.lower()) if len(w) > 3)
            content_lower = truncated.lower()
            if query_words:
                overlap = sum(1 for w in query_words if w in content_lower)
                relevance_ratio = overlap / len(query_words)
            else:
                relevance_ratio = 1.0
            confidence = round(0.5 + 0.43 * min(relevance_ratio, 1.0), 2)

            return ToolResult(
                content=truncated,
                source=f"Document: {matched.name}",
                confidence=confidence, #0.93,
                metadata={
                    "file": str(matched.resolve()),
                    "size_chars": len(content),
                    "extension": matched.suffix,
                    "full_content": content,
                },
            )

        except Exception as e:
            return ToolResult(
                content=f"Error reading '{matched.name}': {str(e)}",
                source=f"Document: {matched.name}",
                confidence=0.3,
            )

    def _find_best_match(self, query: str, files: list) -> Optional[Path]:
        """
        Fallback matching used only when no explicit file_path is given
        (e.g. when document_reader is called as a tool inside the agentic
        loop rather than via researcher.py's pre-check). Deterministic,
        content-aware: checks actual file content for query keyword
        overlap, not just the filename, since filenames are often
        uninformative (e.g. "document (26) (1).pdf").
        """
        q = query.lower().replace('_', ' ').replace('-', ' ')
        query_words = {
            w for w in re.findall(r'\w+', q)
            if len(w) > 3 and w not in _GENERIC_STOPWORDS
        }

        if not query_words:
            return files[0] if len(files) == 1 else None

        best_file, best_score = None, 0

        #if best_score < 3 and len(files) > 1:
        #    return None

        for f in files:
            #name_words = set(re.findall(r'\w+', f.stem.lower()))
            name_words = set(re.findall(r'\w+', f.stem.lower().replace('_', ' ').replace('-', ' ')))
            score = len(query_words & name_words) * 3  # filename overlap weighted higher

            # Cheap content check: read a small prefix and count keyword hits,
            # without fully extracting the whole file just to rank candidates.
            try:
                preview = self._cheap_preview(f)
                preview_lower = preview.lower()
                score += sum(1 for w in query_words if w in preview_lower)
            except Exception:
                pass

            if score > best_score:
                best_score = score
                best_file = f

        if best_score >= 3:
            return best_file

        if len(files) == 1:
            return files[0]

        if best_file is not None and best_score > 0:
            return best_file

        return None

    def _cheap_preview(self, path: Path, max_chars: int = 6000) -> str:
        """Fast, best-effort preview for ranking candidates in
        _find_best_match — not the full extraction path."""
        ext = path.suffix.lower()
        try:
            if ext in (".txt", ".md", ".csv"):
                return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
            elif ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                text = ""
                for page in reader.pages[:12]:
                    text += page.extract_text() or ""
                    if len(text) >= max_chars:
                        break
                return text[:max_chars]
            elif ext in (".docx", ".doc"):
                import docx
                doc = docx.Document(str(path))
                text = "\n".join(p.text for p in doc.paragraphs[:200])
                return text[:max_chars]
        except Exception:
            return ""
        return ""

    def quick_preview(self, path: Path, max_pages: int = 3, dpi: int = 100) -> str:
        """
        Fast, low-fidelity preview for relevance SCORING only — not used to
        answer questions. For text-based PDFs, just extracts the first few
        pages normally (cheap already). For scanned PDFs, OCRs only the
        first `max_pages` at reduced DPI instead of the whole document, so
        ranking candidate documents doesn't cost a full OCR pass.
        """
        ext = path.suffix.lower()
        if ext != ".pdf":
            return self._cheap_preview(path)

        try:
            reader = PdfReader(str(path))
            text_parts = []
            for i, page in enumerate(reader.pages[:max_pages]):
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
            text = "\n".join(text_parts)
            if len(text.strip()) > 200:
                return text  # real text layer exists, no OCR needed

            # Scanned PDF — OCR just a few pages, at lower DPI, no cache write.
            pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            poppler_path = r"C:\Release-26.02.0-0\poppler-26.02.0\Library\bin"

            images = convert_from_path(
                str(path), first_page=1, last_page=min(max_pages, len(reader.pages)),
                dpi=dpi, fmt="jpeg", poppler_path=poppler_path,
            )
            ocr_text = " ".join(
                pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                for img in images
            )
            return ocr_text
        except Exception:
            return ""

    def _read_file(self, path: Path):
        """Returns (content, truncation_note)."""
        ext = path.suffix.lower()

        if ext in (".txt", ".md", ".csv"):
            return self._read_text(path), None
        elif ext == ".pdf":
            return self._read_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._read_docx(path), None
        else:
            return self._read_text(path), None

    def _read_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1", errors="ignore")

    def _read_pdf(self, path: Path) -> Tuple[str, Optional[str]]:
        """Read PDF using pypdf, bounded by MAX_PDF_PAGES so an extremely
        long document fails fast with a clear note instead of taking
        minutes and looking like a hang."""
        cache_path = path.parent / f"{path.name}.ocr_cache.json"

        if cache_path.exists():
            try:
                if cache_path.stat().st_mtime >= path.stat().st_mtime:
                    cached = json.loads(cache_path.read_text(encoding="utf-8"))
                    print(f"      [DEBUG] OCR cache HIT for '{path.name}' — skipping OCR.")
                    return cached["content"], cached.get("note")
                else:
                    print(f"      [DEBUG] OCR cache exists for '{path.name}' but is stale (source file newer).")

            except Exception as e:
                print(f"      [DEBUG] OCR cache exists but failed to read: {e}")
                pass  # corrupt/unreadable cache — fall through and regenerate
        else:
            print(f"      [DEBUG] No OCR cache found for '{path.name}' at {cache_path}")

        try:
            reader = PdfReader(str(path))
            total_pages = len(reader.pages)
            pages_to_read = min(total_pages, MAX_PDF_PAGES)

            text_parts = []
            for i in range(pages_to_read):
                page_text = reader.pages[i].extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {i + 1}]\n{page_text}")

            content = "\n\n".join(text_parts) if text_parts else ""
            total_text_len = len(content)
            
            # If we got decent text → return it
            if total_text_len > 400:
                note = None
                if total_pages > MAX_PDF_PAGES:
                    note = (
                        f"This PDF has {total_pages} pages; only the first "
                        f"{MAX_PDF_PAGES} were processed for performance reasons. "
                        f"If the answer is likely in a later section, consider "
                        f"splitting that section into its own file."
                    )

                return content, note

            # ---------- OCR FALLBACK for scanned PDFs ----------
            print(f"📄 '{path.name}' appears scanned → starting OCR...")

            try:

                # Force paths (adjust if your install locations are different)
                pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
                poppler_path = r"C:\Release-26.02.0-0\poppler-26.02.0\Library\bin"

                # Adaptive page limit
                file_size_mb = path.stat().st_size / (1024 * 1024)
                
                if file_size_mb > 100:
                    last_page = 60
                elif file_size_mb > 40:
                    last_page = 45
                elif file_size_mb > 15:
                    last_page = 30
                else:
                    last_page = 15

                last_page = min(last_page, total_pages)

                images = convert_from_path(
                    str(path),
                    first_page=1,
                    last_page=last_page,
                    dpi=140,
                    fmt="jpeg",
                    thread_count=3,
                    poppler_path=poppler_path,
                )

                def _ocr_one_page(args):
                    i, img = args
                    text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                    return i, text

                max_workers = min(4, os.cpu_count() or 2)
                print(f"   → OCR-ing {len(images)} pages using {max_workers} parallel workers...")


                results = {}
                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [executor.submit(_ocr_one_page, (i, img)) for i, img in enumerate(images)]
                    completed = 0
                    for future in concurrent.futures.as_completed(futures):
                        i, text = future.result()
                        results[i] = text
                        completed += 1
                        print(f"   → OCR page {completed}/{len(images)} done (page #{i+1})")

                ocr_parts = []
                for i in range(len(images)):
                    text = results.get(i, "")
                    if text and len(text.strip()) > 40:
                        ocr_parts.append(f"[Page {i+1} - OCR]\n{text.strip()}")

                """
                ocr_parts = []
                for i, img in enumerate(images):
                    print(f"   → OCR page {i+1}/{len(images)}...")
                    text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                    if text and len(text.strip()) > 40:
                        ocr_parts.append(f"[Page {i+1} - OCR]\n{text.strip()}")
                 """
                if ocr_parts:
                    result = "\n\n".join(ocr_parts)
                    print(f"✅ OCR completed. Extracted {len(result)} characters.")
                    note = f"OCR used on first {last_page} pages (scanned PDF)."

                    try:
                        cache_path.write_text(
                            json.dumps({"content": result, "note": note}),
                            encoding="utf-8",
                         )
                        print(f"      [DEBUG] OCR cache WRITTEN to {cache_path}")
                    except Exception as e:
                        print(f"      [DEBUG] ⚠️ OCR cache write FAILED: {e}")
                        pass  # caching is an optimization — never let it break a real read

                    return result, note
                else:
                    return (
                        f"WARNING: OCR produced little/no text from '{path.name}'.",
                        "OCR failed to extract usable text"
                    )

            except Exception as ocr_err:
                print(f"❌ OCR failed: {ocr_err}")
                return (
                    f"PDF appears scanned but OCR failed: {ocr_err}",
                    "OCR error"
                )    
                
        except ImportError:
            return (
                "pypdf not installed. Run:\n"
                "  pip install pypdf\n"
                "Then restart and try again."
            ), None
        
        except Exception as e:
            return f"PDF read error: {str(e)}", None

    def _read_docx(self, path: Path) -> str:
        try:
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            return (
                "python-docx not installed. Run:\n"
                "  pip install python-docx\n"
                "Then restart and try again."
            )
        except Exception as e:
            return f"DOCX read error: {str(e)}"

    def list_documents(self) -> list:
        return [
            {
                "name": f.name,
                "type": f.suffix.upper().lstrip("."),
                "size_kb": round(f.stat().st_size / 1024, 1),
                "path": str(f.resolve()),
            }
            for f in self.docs_folder.rglob("*")
            if f.is_file() and f.suffix.lower() in SUPPORTED
        ]
